"""
DOCX Document Parser
"""
from pathlib import Path
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX files"""
    
    @staticmethod
    def parse(file_path: Path) -> Dict[str, Any]:
        """
        Parse DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with 'content' and 'metadata'
        """
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        tables_text.append(row_text)
            
            # Combine content
            content_parts = []
            if paragraphs:
                content_parts.append('\n\n'.join(paragraphs))
            if tables_text:
                content_parts.append('\n'.join(tables_text))
            
            content = '\n\n---\n\n'.join(content_parts)
            
            # Metadata
            word_count = len(content.split())
            char_count = len(content)
            
            metadata = {
                'format': 'docx',
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'word_count': word_count,
                'character_count': char_count
            }
            
            # Core properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': str(props.created) if props.created else '',
                })
            
            logger.info(f"Parsed DOCX: {len(paragraphs)} paragraphs, {word_count} words")
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise
