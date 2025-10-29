# document_processing/processors/general_processor.py
"""
المعالج العام للمستندات - من aisearchmm
يدعم PDF, DOCX, TXT وصيغ أخرى
"""

from typing import Optional
from pathlib import Path

from .base_processor import BaseDocumentProcessor, ProcessedDocument, DocumentMetadata
from document_processing.parsers.pdf_parser import PDFParser
from document_processing.chunking.multilingual_splitter import MultilingualTextSplitter
from core.config import config
from core.exceptions import DocumentProcessingError
from utilities.logger import logger


class GeneralDocumentProcessor(BaseDocumentProcessor):
    """معالج المستندات العامة"""
    
    def __init__(self):
        super().__init__()
        
        # الصيغ المدعومة
        self.supported_formats = ["pdf", "txt", "md", "docx", "doc"]
        
        # المحللات
        self.pdf_parser = PDFParser()
        
        # مقسم النصوص
        self.text_splitter = MultilingualTextSplitter(
            chunk_size=config.document_processing.chunk_size,
            chunk_overlap=config.document_processing.chunk_overlap
        )
        
        logger.info(f"Initialized {self.processor_name} with formats: {self.supported_formats}")
    
    async def process(self, file_path: str) -> ProcessedDocument:
        """
        معالجة مستند عام
        
        Args:
            file_path: مسار الملف
            
        Returns:
            ProcessedDocument: المستند المعالج
        """
        logger.info(f"Processing document: {file_path}")
        
        # التحقق من الملف
        self._validate_file(file_path)
        
        # التحقق من إمكانية المعالجة
        if not self.can_process(file_path):
            raise DocumentProcessingError(
                f"Unsupported file format for {file_path}",
                {"file_path": file_path, "supported": self.supported_formats}
            )
        
        # استخراج البيانات الوصفية
        metadata = self._extract_metadata(file_path)
        
        # استخراج المحتوى حسب النوع
        file_type = self._get_file_type(file_path)
        
        if file_type == "pdf":
            content_data = await self._process_pdf(file_path)
        elif file_type in ["txt", "md"]:
            content_data = await self._process_text(file_path)
        elif file_type in ["docx", "doc"]:
            content_data = await self._process_docx(file_path)
        else:
            raise DocumentProcessingError(
                f"No processor available for {file_type}",
                {"file_type": file_type}
            )
        
        # تحديث البيانات الوصفية
        metadata.language = self._detect_language(content_data["text"])
        if "page_count" in content_data:
            metadata.page_count = content_data["page_count"]
        if "metadata" in content_data:
            metadata.author = content_data["metadata"].get("author")
            metadata.title = content_data["metadata"].get("title")
        
        # تقسيم النص
        chunks = self.text_splitter.split_text(content_data["text"])
        
        # إنشاء المستند المعالج
        doc = ProcessedDocument(
            id=self._generate_document_id(file_path),
            content=content_data["text"],
            metadata=metadata,
            chunks=chunks,
            tables=content_data.get("tables", []),
            images=content_data.get("images", [])
        )
        
        logger.info(
            f"✅ Document processed successfully: "
            f"{len(doc.content)} chars, {len(doc.chunks)} chunks, "
            f"{metadata.language} language"
        )
        
        return doc
    
    async def _process_pdf(self, file_path: str) -> dict:
        """معالجة PDF"""
        logger.debug(f"Processing PDF: {file_path}")
        result = self.pdf_parser.parse(file_path)
        return result
    
    async def _process_text(self, file_path: str) -> dict:
        """معالجة ملفات TXT/MD"""
        logger.debug(f"Processing text file: {file_path}")
        
        try:
            # قراءة المحتوى
            text = Path(file_path).read_text(encoding='utf-8')
            
            return {
                "text": text,
                "page_count": 1
            }
        
        except UnicodeDecodeError:
            # محاولة ترميزات أخرى
            for encoding in ['latin-1', 'cp1256', 'iso-8859-1']:
                try:
                    text = Path(file_path).read_text(encoding=encoding)
                    logger.warning(f"Used {encoding} encoding for {file_path}")
                    return {"text": text, "page_count": 1}
                except:
                    continue
            
            raise DocumentProcessingError(
                f"Failed to decode text file: {file_path}",
                {"file_path": file_path}
            )
    
    async def _process_docx(self, file_path: str) -> dict:
        """معالجة DOCX"""
        logger.debug(f"Processing DOCX: {file_path}")
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # استخراج النص من الفقرات
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # استخراج النص من الجداول
            tables_text = []
            for table in doc.tables:
                table_text = "\n".join([
                    "\t".join([cell.text for cell in row.cells])
                    for row in table.rows
                ])
                tables_text.append(table_text)
            
            # دمج النص
            full_text = text
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)
            
            return {
                "text": full_text,
                "page_count": len(doc.sections),
                "tables": [{"content": t} for t in tables_text]
            }
        
        except ImportError:
            logger.warning("python-docx not installed, reading DOCX as text")
            # محاولة قراءة كنص عادي (سيكون محدوداً)
            return await self._process_text(file_path)
        
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to process DOCX: {str(e)}",
                {"file_path": file_path}
            )